"""
Generate ICDF project reference document (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

OUTPUT = r"c:\final yr project\Cyber-Defense-Framework\ICDF_Project_Reference.docx"


def set_heading_style(doc):
    styles = doc.styles
    for i in range(1, 4):
        name = f"Heading {i}"
        if name in styles:
            styles[name].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Intelligent Cyber Defense Framework\n"
        "Project Reference & Improvement Guide"
    )
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "Adaptive Network Security using Reinforcement Learning\n"
        "and Cyber Threat Intelligence"
    )
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nGenerated: {datetime.now().strftime('%d %B %Y')}\n").font.size = Pt(11)
    meta.add_run("Final-Year Undergraduate Project Reference Document").font.size = Pt(11)

    doc.add_page_break()


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def build_document():
    doc = Document()
    set_heading_style(doc)
    add_title_page(doc)

    # PART 1: COMPREHENSIVE DESCRIPTION
    doc.add_heading("Part 1: Comprehensive Project Description", level=1)

    doc.add_heading("1. What Is This Project?", level=2)
    doc.add_paragraph(
        "The Intelligent Cyber Defense Framework (ICDF) is an AI-powered adaptive network "
        "security system developed as a final-year cybersecurity project. It ingests network "
        "traffic features and an IP address, detects attacks using machine learning, enriches "
        "findings with external Cyber Threat Intelligence (CTI), calculates an overall risk "
        "score, and uses Reinforcement Learning (RL) to recommend an appropriate defensive action."
    )
    doc.add_paragraph(
        "Unlike traditional Intrusion Detection Systems (IDS) that primarily alert and rely on "
        "manual intervention, ICDF chains detection, intelligence gathering, risk assessment, "
        "and automated decision support into a single end-to-end pipeline exposed via a REST API "
        "and visualized through a React-based Security Operations Center (SOC) dashboard."
    )

    doc.add_heading("2. Problem Statement", level=2)
    for item in [
        "Traditional IDS tools detect attacks but do not decide what to do next.",
        "Alerts often lack real-world reputation context from threat intelligence feeds.",
        "All alerts are treated similarly without a unified, weighted risk score.",
        "Incident response depends heavily on manual SOC playbooks.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph(
        "ICDF addresses these gaps by combining Machine Learning detection, VirusTotal and "
        "AbuseIPDB CTI enrichment, a custom Risk Engine, and a Deep Q-Network (DQN) for "
        "adaptive response recommendations."
    )

    doc.add_heading("3. High-Level Architecture", level=2)
    doc.add_paragraph(
        "Network Features + IP Address\n"
        "    → Random Forest Attack Detection\n"
        "    → VirusTotal + AbuseIPDB (CTI)\n"
        "    → Risk Engine\n"
        "    → Deep Q-Network (DQN)\n"
        "    → Decision Engine / Response Actions\n"
        "    → Flask REST API\n"
        "    → MongoDB + React SOC Dashboard"
    )

    doc.add_heading("4. End-to-End Pipeline", level=2)

    doc.add_heading("4.1 Input", level=3)
    doc.add_paragraph(
        "The system accepts an IP address and exactly 78 numeric features representing a "
        "CICIDS2017-style network traffic feature vector. These are submitted via POST /analyze "
        "or through the React Live Analyze page."
    )

    doc.add_heading("4.2 Random Forest Attack Detection", level=3)
    doc.add_paragraph(
        "A trained Random Forest classifier (random_forest_model.pkl) predicts attack type, "
        "confidence (0–100%), and severity (0–100) mapped per attack class."
    )
    doc.add_paragraph("Supported attack classes:")
    attacks = (
        "BENIGN, Bot, PortScan, FTP-Patator, SSH-Patator, DoS Hulk, DoS GoldenEye, "
        "DoS Slowhttptest, DoS slowloris, DDoS, Heartbleed, Infiltration, "
        "Web Attack – Brute Force, Web Attack – Sql Injection, Web Attack – XSS"
    )
    doc.add_paragraph(attacks)

    doc.add_heading("4.3 Cyber Threat Intelligence (CTI)", level=3)
    doc.add_paragraph("VirusTotal returns: malicious, suspicious, harmless, undetected counts.")
    doc.add_paragraph(
        "AbuseIPDB returns: abuse confidence, country, usage type, total reports, whitelist status."
    )
    doc.add_paragraph(
        "VirusTotal scores are normalized to 0–100 for risk calculation. API failures are "
        "handled gracefully without crashing the pipeline."
    )

    doc.add_heading("4.4 Risk Engine", level=3)
    add_table(
        doc,
        ["Factor", "Weight"],
        [
            ["Attack Severity", "40%"],
            ["Model Confidence", "20%"],
            ["VirusTotal Score", "20%"],
            ["AbuseIPDB Score", "20%"],
        ],
    )
    doc.add_paragraph("Risk levels: SAFE, LOW, MEDIUM, HIGH, CRITICAL")

    doc.add_heading("4.5 Reinforcement Learning (DQN)", level=3)
    add_table(
        doc,
        ["Action Code", "Meaning"],
        [
            ["NO_ACTION", "Allow Traffic"],
            ["ALERT_ADMIN", "Alert Administrator"],
            ["BLOCK_IP", "Block IP"],
            ["ISOLATE_HOST", "Isolate Host"],
        ],
    )
    doc.add_paragraph(
        "State: (attack severity, risk score). Trained with PyTorch using a custom environment "
        "and reward function. Model saved as dqn_model.pth."
    )

    doc.add_heading("4.6 Decision Engine & Response Actions", level=3)
    doc.add_paragraph(
        "The Decision Engine loads the DQN, predicts an action, and triggers the corresponding "
        "response. Actions are simulated via log files under logs/ (not real firewall enforcement)."
    )

    doc.add_heading("4.7 Flask REST API", level=3)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/health", "GET", "API health check"],
            ["/analyze", "POST", "Run full analysis pipeline"],
            ["/history", "GET", "List saved analyses (limit, skip)"],
            ["/history/<id>", "GET", "Fetch single analysis"],
        ],
    )

    doc.add_heading("4.8 MongoDB Storage", level=3)
    doc.add_paragraph(
        "Analysis results stored in intelligent_cyber_defense.analysis_history. MongoDB is "
        "optional for viewing the UI; Live Analyze works without it. History and analytics "
        "require MongoDB."
    )

    doc.add_heading("4.9 React SOC Dashboard", level=3)
    for page in [
        "Overview — KPIs, charts, recent threats",
        "Live Analyze — submit IP + 78 features with demo generator",
        "Threat History — searchable, filterable table",
        "Detail — full analysis breakdown",
        "Blocked / Isolated / Alerts — filtered by RL action",
        "Analytics — trends, distributions, scatter plots",
    ]:
        add_bullet(doc, page)
    doc.add_paragraph("Frontend: http://localhost:3000 | Backend: http://localhost:5000")

    doc.add_heading("5. Machine Learning Pipeline (Offline Training)", level=2)
    for step in [
        "Preprocessing — merge CICIDS2017 raw CSV files",
        "Data cleaning — deduplication, infinity handling, balanced sampling",
        "Feature engineering — label encoding, StandardScaler, train/test split",
        "Random Forest training — saved as .pkl",
        "DQN training — custom RL environment with experience replay",
    ]:
        add_bullet(doc, step)

    doc.add_heading("6. Technology Stack", level=2)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Language", "Python"],
            ["ML", "scikit-learn (Random Forest)"],
            ["RL", "PyTorch (Deep Q-Network)"],
            ["CTI", "VirusTotal API, AbuseIPDB API"],
            ["Backend", "Flask, Flask-CORS"],
            ["Database", "MongoDB (PyMongo)"],
            ["Frontend", "React, Axios, Chart.js, React Router"],
            ["Dataset", "CICIDS2017"],
        ],
    )

    doc.add_heading("7. Example API Response", level=2)
    doc.add_paragraph(
        'POST /analyze returns: ip_address, prediction (attack, severity, confidence), '
        "virustotal, abuseipdb, risk (risk_score, risk_level), decision (action, status), "
        "analysis_id, saved_to_mongodb."
    )

    doc.add_heading("8. What Makes It Intelligent", level=2)
    for item in [
        "Detection — ML classifies attack type from network features",
        "Context — CTI adds real-world reputation beyond the model",
        "Prioritization — Risk engine ranks threats by severity and intelligence",
        "Adaptation — RL selects response based on severity + risk, not fixed rules only",
        "Automation — One API call runs the entire chain",
        "Observability — Dashboard and MongoDB history for SOC-style review",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9. Intended Use Cases", level=2)
    for item in [
        "Enterprise / campus network security demonstrations",
        "SOC-style monitoring and triage",
        "Academic research on ML + CTI + RL for cyber defense",
        "Proof-of-concept for automated incident response",
    ]:
        add_bullet(doc, item)

    doc.add_heading("10. Current Limitations", level=2)
    for item in [
        "No live packet capture — uses 78 pre-extracted features, not real-time PCAP",
        "Simulated defense — block/isolate/alert are logged, not enforced on network gear",
        "CTI dependency — VirusTotal/AbuseIPDB require API keys and have rate limits",
        "MongoDB optional — UI works without it; history features are limited",
        "Academic scope — designed for demonstration, not production deployment as-is",
    ]:
        add_bullet(doc, item)

    doc.add_heading("11. Project Structure", level=2)
    doc.add_paragraph(
        "Cyber-Defense-Framework/\n"
        "  ml/          — ML, RL, risk engine, decision logic\n"
        "  cti/         — VirusTotal & AbuseIPDB clients\n"
        "  backend/     — Flask API, pipeline, MongoDB\n"
        "  frontend/    — React SOC dashboard\n"
        "  logs/        — Simulated response logs\n"
        "  dataset/     — CICIDS2017 (gitignored)"
    )

    doc.add_heading("12. One-Paragraph Summary (for Reports / Viva)", level=2)
    doc.add_paragraph(
        "The Intelligent Cyber Defense Framework is an adaptive network security system that "
        "ingests network traffic features and an IP address, detects attacks using a Random Forest "
        "model trained on CICIDS2017, enriches results with VirusTotal and AbuseIPDB threat "
        "intelligence, computes a weighted risk score, and uses a Deep Q-Network to recommend "
        "defensive actions such as allowing traffic, alerting an administrator, blocking an IP, "
        "or isolating a host. The full pipeline is exposed through a Flask REST API, results are "
        "stored in MongoDB, and a React-based SOC dashboard provides visualization, live analysis, "
        "threat history, and security analytics—forming an end-to-end AI-driven cyber defense "
        "platform suitable for automated incident response research and demonstration."
    )

    doc.add_page_break()

    # PART 2: ASSESSMENT & IMPROVEMENTS
    doc.add_heading("Part 2: Project Assessment & Recommended Improvements", level=1)

    doc.add_heading("Overall Assessment", level=2)
    doc.add_paragraph(
        "For a final-year undergraduate project, ICDF is significantly above average. It "
        "combines multiple AI techniques (ML + RL), external threat intelligence, a backend API, "
        "a database, and a modern frontend. Many final-year cybersecurity projects stop at "
        "'intrusion detection using machine learning.' This project extends to detection → "
        "enrichment → risk assessment → automated decision support → visualization."
    )

    doc.add_heading("Rating by Area", level=2)
    add_table(
        doc,
        ["Area", "Rating"],
        [
            ["Problem relevance", "10/10"],
            ["Machine Learning integration", "9.5/10"],
            ["Reinforcement Learning", "9/10"],
            ["Backend architecture", "9/10"],
            ["Frontend potential", "9.5/10 (after SOC dashboard)"],
            ["Innovation", "9/10"],
            ["Academic value", "9.5/10"],
        ],
    )

    doc.add_heading("Recommended Improvements (Priority Order)", level=2)

    improvements = [
        (
            "1. Make the Dashboard Feel Like a Real SOC (Highest Priority)",
            "★★★★★",
            [
                "Live threat feed and attack timeline",
                "Colored severity badges and attack heatmaps",
                "Recent incidents panel",
                "Animated risk gauge",
                "Interactive charts with search/filtering",
                "API health indicator",
                "Loading pipeline animation during analysis",
                "Gives the impression of an enterprise security platform",
            ],
        ),
        (
            "2. Explain WHY the RL Chose an Action",
            "★★★★★",
            [
                "Instead of only 'Risk = HIGH, Action = BLOCK_IP', show reasoning:",
                "Attack classified as DDoS",
                "Severity = 95, Model confidence = 96%",
                "VirusTotal detected malicious activity",
                "Overall Risk Score = 72.5",
                "RL recommends blocking because high-risk attacks receive higher reward when blocked during training",
            ],
        ),
        (
            "3. Risk Breakdown Visualization",
            "★★★★★",
            [
                "Show weighted components: Attack Severity 40%, Model Confidence 20%, VirusTotal 20%, AbuseIPDB 20%",
                "Use stacked progress bar for visual impact",
            ],
        ),
        (
            "4. Threat Intelligence Cards",
            "★★★★",
            [
                "VirusTotal card: Malicious, Suspicious, Harmless, Undetected counts",
                "AbuseIPDB card: Country, Confidence, Reports, Usage Type",
                "Replace raw JSON with SOC-style presentation",
            ],
        ),
        (
            "5. Add Explainability (Random Forest Feature Importance)",
            "★★★★",
            [
                "Display top important features (global feature importance)",
                "Examples: Flow Duration, Packet Length Mean, Fwd Packet Length, Idle Mean, Active Mean",
                "Adds academic value even if not per-prediction SHAP values",
            ],
        ),
        (
            "6. Enhanced Analytics Dashboard",
            "★★★★",
            [
                "Top attack types, countries with highest abuse reports",
                "Risk trends over time, average confidence and severity",
                "Daily/weekly analyses, RL action breakdown, threat level distribution",
                "Splunk-like analytics experience",
            ],
        ),
        (
            "7. Export Reports",
            "★★★★",
            [
                "One-click export: PDF, CSV, JSON",
                "Highly valued by examiners for demonstrations",
            ],
        ),
        (
            "8. Threat Timeline",
            "★★★★",
            [
                "Visual timeline: 10:30 DDoS detected → 10:31 Risk HIGH → 10:31 BLOCK_IP → Stored in MongoDB",
                "More professional than flat history list",
            ],
        ),
        (
            "9. Better Detail Page",
            "★★★★",
            [
                "Split into cards: Prediction, Threat Intelligence, Risk Assessment, Decision, Metadata",
                "Collapsible JSON viewer instead of raw dump",
            ],
        ),
        (
            "10. Pipeline Visualization",
            "★★★★★",
            [
                "Show: Input → RF → VT → AbuseIPDB → Risk → DQN → Decision → MongoDB",
                "Highlight each stage while loading",
                "Very impressive during live demo",
            ],
        ),
        (
            "11. Interactive Architecture Page",
            "★★★★",
            [
                "Clickable architecture diagram for viva",
                "Click Random Forest / DQN / Risk Engine for explanations",
                "Excellent for examiner Q&A",
            ],
        ),
        (
            "12. Statistics Cards",
            "★★★",
            [
                "Average Risk, Highest Risk, Blocked count, Alerts count, Safe Traffic %",
            ],
        ),
        (
            "13. Better Search & Filters",
            "★★★",
            [
                "History: search by IP, Attack, Country, Risk, Action",
                "Multiple simultaneous filters",
            ],
        ),
        (
            "14. Consistent Dashboard Theme",
            "★★★",
            [
                "Green = Safe, Yellow = Medium, Orange = High, Red = Critical, Blue = Information",
            ],
        ),
        (
            "15. Animated Risk Gauge",
            "★★★★",
            [
                "Risk meter with progress bar instead of plain number",
                "Shows score + level (e.g. 72.5 HIGH)",
            ],
        ),
        (
            "16. Demo Mode",
            "★★★★★",
            [
                "One button: Generate Demo Attack",
                "Auto-fills random IP + 78 features and runs analysis",
                "Perfect for presentations",
            ],
        ),
        (
            "17. Project Statistics Page",
            "★★★",
            [
                "RF Accuracy, 15 attack classes, 78 features, 4 RL actions",
                "2 CTI sources, MongoDB, Flask, React",
                "Useful reference for examiners",
            ],
        ),
        (
            "18. Improved Error Handling",
            "★★★★",
            [
                "'VirusTotal unavailable. Analysis completed using ML + Risk Engine.'",
                "Instead of raw 401 errors",
            ],
        ),
        (
            "19. Better Visual Icons",
            "★★★",
            [
                "Shield (Prediction), Warning (Risk), Globe (CTI), Brain (RL), Cylinder (Database)",
            ],
        ),
        (
            "20. About the Framework Page",
            "★★★",
            [
                "Explains: Detection → CTI → Risk → Decision → Response",
                "Good for demonstrations and viva",
            ],
        ),
    ]

    for title, impact, items in improvements:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph()
        p.add_run(f"Impact: {impact}").bold = True
        for item in items:
            add_bullet(doc, item)

    doc.add_heading("Optional Advanced Enhancements", level=2)
    for item in [
        "Role-based authentication (admin/analyst login)",
        "Real-time updates via WebSockets or Server-Sent Events",
        "Docker / Docker Compose one-command deployment",
        "Swagger/OpenAPI documentation for Flask API",
        "Unit and integration tests for backend and frontend",
        "Model evaluation page (confusion matrix, precision, recall, F1, ROC)",
    ]:
        add_bullet(doc, item)

    doc.add_heading("What NOT to Add (Out of Scope)", level=2)
    for item in [
        "Live packet capture with Scapy/Wireshark",
        "Real firewall or router automation",
        "SIEM integrations (Splunk, QRadar, Sentinel)",
        "Kubernetes deployment",
        "Distributed microservices",
        "LLM-powered threat analysis",
        "Real SOC ticketing integrations",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph("These are better suited for future work.")

    doc.add_heading("Supervisor-Recommended Priority", level=2)
    for i, item in enumerate(
        [
            "Complete the SOC dashboard with polished UX, charts, and navigation",
            "Add explainability for RF prediction, risk score, and RL recommendation",
            "Enhance analytics with richer charts and summary metrics",
            "Implement export features (CSV/PDF/JSON) and strong error handling",
            "Document architecture and evaluation thoroughly in the project report",
        ],
        1,
    ):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_paragraph(
        "\nWith these improvements, the project presents as a cohesive AI-driven cybersecurity "
        "platform rather than a collection of individual components—a strong and well-rounded "
        "final-year submission."
    )

    doc.add_page_break()

    # PART 3: API & FRONTEND REFERENCE
    doc.add_heading("Part 3: Technical Reference for Development", level=1)

    doc.add_heading("API Contract", level=2)
    doc.add_paragraph("Base URL: http://localhost:5000 (REACT_APP_API_URL)")

    doc.add_heading("POST /analyze", level=3)
    doc.add_paragraph('Body: { "ip_address": "8.8.8.8", "features": [78 floats] }')
    doc.add_paragraph(
        "Response fields: prediction, virustotal, abuseipdb, risk, decision, "
        "analysis_id, saved_to_mongodb, mongodb_error (if Mongo offline)"
    )

    doc.add_heading("GET /history?limit=50&skip=0", level=3)
    doc.add_paragraph("Returns: status, total, count, limit, skip, history[]")

    doc.add_heading("Frontend Routes", level=2)
    add_table(
        doc,
        ["Route", "Page"],
        [
            ["/", "Overview Dashboard"],
            ["/analyze", "Live Analyze"],
            ["/history", "Threat History"],
            ["/history/:id", "Analysis Detail"],
            ["/blocked", "Blocked IPs"],
            ["/isolated", "Isolated Hosts"],
            ["/alerts", "Alerts"],
            ["/analytics", "Analytics"],
        ],
    )

    doc.add_heading("How to Run", level=2)
    doc.add_paragraph("Backend (project root):\n  python -m backend.app")
    doc.add_paragraph("Frontend:\n  cd frontend\n  npm install\n  npm start")
    doc.add_paragraph("MongoDB (optional for history):\n  localhost:27017")
    doc.add_paragraph(
        "API keys in backend/.env:\n  VIRUSTOTAL_API_KEY\n  ABUSEIPDB_API_KEY"
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_document()
